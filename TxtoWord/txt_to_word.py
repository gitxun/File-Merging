import logging
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from tqdm import tqdm


def ensure_rPr(style_element):
    rPr = style_element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style_element.insert(0, rPr)
    return rPr

def set_style_font(style, font_name, font_size_pt, bold, italic, color):
    font = style.font
    font.name = font_name
    font.size = Pt(font_size_pt)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor(*color)

    rPr = ensure_rPr(style.element)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def unify_styles(doc):
    styles_config = {
        'Heading 1': {'font_name': '黑体', 'font_size':16, 'bold': True, 'italic': False, 'color':(0,0,0), 'line_spacing':1.5, 'first_line_indent':0},
        'Heading 2': {'font_name': '黑体', 'font_size':15, 'bold': True, 'italic': False, 'color':(0,0,0), 'line_spacing':1.5, 'first_line_indent':0},
        'Heading 3': {'font_name': '黑体', 'font_size':14, 'bold': True, 'italic': False, 'color':(0,0,0), 'line_spacing':1.5, 'first_line_indent':0},
        'Heading 4': {'font_name': '黑体', 'font_size':12, 'bold': True, 'italic': False, 'color':(0,0,0), 'line_spacing':1.25, 'first_line_indent':0},
        'Heading 5': {'font_name': '黑体', 'font_size':12, 'bold': True, 'italic': False, 'color':(0,0,0), 'line_spacing':1.25, 'first_line_indent':0},
        'Normal':    {'font_name': '宋体', 'font_size':12, 'bold': False, 'italic': False, 'color':(0,0,0), 'line_spacing':1.25, 'first_line_indent':0.74},
    }

    for style_name, cfg in styles_config.items():
        if style_name not in doc.styles:
            logging.info(f"样式{style_name}不存在，跳过")
            continue
        style = doc.styles[style_name]
        set_style_font(style, cfg['font_name'], cfg['font_size'], cfg['bold'], cfg['italic'], cfg['color'])

        if style.paragraph_format:
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            style.paragraph_format.line_spacing = cfg['line_spacing']
            style.paragraph_format.first_line_indent = Cm(cfg['first_line_indent'])

def apply_style_properties(paragraph, style_config):
    """直接应用样式配置到段落，不依赖于样式名称"""
    # 应用字体设置到所有runs
    for run in paragraph.runs:
        font = run.font
        font.name = style_config['font_name']
        font.size = Pt(style_config['font_size'])
        font.bold = style_config['bold']
        font.italic = style_config['italic']
        font.color.rgb = RGBColor(*style_config['color'])
        
        # 设置东亚字体 (中文)
        r = run._element.get_or_add_rPr()
        rFonts = r.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), style_config['font_name'])
    
    # 应用段落格式
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = style_config['line_spacing']
    paragraph.paragraph_format.first_line_indent = Cm(style_config['first_line_indent'])

def copy_paragraph(source_paragraph, target_document, force_style=None):
    p = target_document.add_paragraph()
    
    style_name = force_style if force_style else source_paragraph.style.name
    # 判断目标文档是否有该样式
    if style_name not in target_document.styles:
        # 如果没有，改用默认样式
        style_name = 'Normal'
    
    p.style = style_name

    for run in source_paragraph.runs:
        r = p.add_run(run.text)
        r.bold = run.bold
        r.italic = run.italic
        r.underline = run.underline
        if run.font.color.rgb:
            r.font.color.rgb = run.font.color.rgb
        if run.font.size:
            r.font.size = run.font.size

    p.alignment = source_paragraph.alignment
    p.paragraph_format.line_spacing = source_paragraph.paragraph_format.line_spacing
    p.paragraph_format.first_line_indent = source_paragraph.paragraph_format.first_line_indent
    
    # 返回创建的段落，以便后续处理
    return p

def process_and_merge_all_files(folder_path, merged_filename="merged.docx"):
    word_folder = os.path.join(folder_path, "merging_files", "word_files")
    output_path = os.path.join(word_folder, merged_filename)

    # 如果合并文档已存在，先删除
    if os.path.exists(output_path):
        try:
            os.remove(output_path)
            logging.info(f"已删除已存在的合并文档：{merged_filename}")
        except Exception as e:
            logging.error(f"删除合并文档时出错：{e}")
    
    # 如果 formated.docx 存在，也删除
    formated_path = os.path.join(word_folder, "formated.docx")
    if os.path.exists(formated_path):
        try:
            os.remove(formated_path)
            logging.info(f"已删除已存在的格式化文档：formated.docx")
        except Exception as e:
            logging.error(f"删除格式化文档时出错：{e}")

    merged_doc = Document()
    
    # 统一合并文档的样式
    unify_styles(merged_doc)

    styles_config = {
        'Heading 1': {'font_name': '黑体', 'font_size':16, 'bold': True, 'italic': False, 'color':(0,0,0), 'line_spacing':1.5, 'first_line_indent':0},
        'Heading 2': {'font_name': '黑体', 'font_size':15, 'bold': True, 'italic': False, 'color':(0,0,0), 'line_spacing':1.5, 'first_line_indent':0},
        'Heading 3': {'font_name': '黑体', 'font_size':14, 'bold': True, 'italic': False, 'color':(0,0,0), 'line_spacing':1.5, 'first_line_indent':0},
        'Heading 4': {'font_name': '黑体', 'font_size':12, 'bold': True, 'italic': False, 'color':(0,0,0), 'line_spacing':1.25, 'first_line_indent':0},
        'Heading 5': {'font_name': '黑体', 'font_size':12, 'bold': True, 'italic': False, 'color':(0,0,0), 'line_spacing':1.25, 'first_line_indent':0},
        'Normal':    {'font_name': '宋体', 'font_size':12, 'bold': False, 'italic': False, 'color':(0,0,0), 'line_spacing':1.25, 'first_line_indent':0.74},
    }

    # 获取所有需要合并的 docx 文件列表，并排除合并输出文件
    file_list = [
        fname for fname in sorted(os.listdir(word_folder))
        if fname.lower().endswith('.docx') and fname != merged_filename
    ]

    # tqdm 包裹 file_list，即可显示进度条
    for fname in tqdm(file_list, desc="合并进度"):
        fpath = os.path.join(word_folder, fname)

        # 打开子文档
        sub_doc = Document(fpath)

        # 复制内容到合并文档
        for para in sub_doc.paragraphs:
            if para.text.strip():
                # 确定段落的样式名称
                style_name = para.style.name if para.style else 'Normal'
                
                # 复制段落并保留基本格式
                new_para = copy_paragraph(para, merged_doc, force_style=style_name)
                
                # 根据样式名称直接应用正确的格式
                if style_name in styles_config:
                    apply_style_properties(new_para, styles_config[style_name])
                else:
                    # 如果不是预定义的样式，使用Normal样式
                    apply_style_properties(new_para, styles_config['Normal'])

        # 添加分页符
        merged_doc.add_page_break()
        logging.info(f"已处理并合并：{fname}")

    # 保存合并文档
    merged_doc.save(output_path)
    logging.info(f"合并完成，保存为 {merged_filename}")
    return output_path

# 使用示例
if __name__ == "__main__":
    root_dir = r"D:\python_workspace\LLM_apply\FileMerge\projects_txt_modules"  # 这里替换成您的根目录路径
    try:
        merged_file_path = process_and_merge_all_files(root_dir)
        if merged_file_path:
            logging.info(f"合并文档生成成功：{merged_file_path}")
    except Exception as e:
        logging.info(f"出错：{e}")
