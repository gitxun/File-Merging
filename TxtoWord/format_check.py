import os
import re
from docx import Document

def remove_inner_empty_lines(doc):
    """
    去除段落内部的空行（换行符导致的空白行）
    """
    for para in doc.paragraphs:
        lines = para.text.splitlines()
        non_empty_lines = [line for line in lines if line.strip()]
        para.text = '\n'.join(non_empty_lines)


def correct_and_convert_numbered_paragraphs(root_folder_path):
    # 构建目标文件路径 merged.docx
    docx_path = os.path.join(root_folder_path, 'merging_files', 'word_files', 'merged.docx')
    if not os.path.isfile(docx_path):
        print(f"错误：找不到文件 {docx_path}")
        return

    doc = Document(docx_path)

    # 删除段落内部空行
    remove_inner_empty_lines(doc)

    for para in doc.paragraphs:
        style_name = para.style.name

        if '标题' in style_name or 'Heading' in style_name:
            continue

        text = para.text.strip()
        if not text:
            continue

        # 处理“- ”开头的无序列表转编号
        if text.count('- ') > 0:
            first_dash_pos = text.find('- ')
            lead_text = text[:first_dash_pos].strip()
            list_text = text[first_dash_pos:]

            parts = [p.strip() for p in list_text.split('- ') if p.strip()]
            new_items = []
            for i, part in enumerate(parts, start=1):
                part_clean = re.sub(r'[：:；;。.]$', '', part)
                if i == len(parts):
                    part_clean += '。'
                else:
                    part_clean += '；'
                new_items.append(f"  {i}) {part_clean}")

            if lead_text:
                para.text = f"{lead_text}\n" + '\n'.join(new_items)
            else:
                para.text = '\n'.join(new_items)
            continue

        # 处理无编号但多行且每行含“：”的段落
        if ('\n' in para.text or (para.text.count('：') > 0 and para.text.count('\n') > 0)):
            lines = [line.strip() for line in para.text.splitlines() if line.strip()]
            if len(lines) > 1 and all('：' in line for line in lines):
                new_lines = []
                for i, line in enumerate(lines):
                    line_clean = re.sub(r'[：:；;。.]$', '', line)
                    if i == len(lines) - 1:
                        line_clean += '。'
                    else:
                        line_clean += '；'
                    new_lines.append(f"  {line_clean}")
                para.text = '\n'.join(new_lines)
                continue

    # 输出路径同样放在 word_files 目录下，命名为 formated.docx
    output_path = os.path.join(root_folder_path, 'merging_files', 'word_files', 'formated.docx')
    doc.save(output_path)
    print(f"处理完成，文件已保存至 {output_path}")
    return output_path   # 返回文件路径


if __name__ == "__main__":
    folder_path = r'D:\python_workspace\LLM_apply\FileMerge\projects_txt_modules'
    correct_and_convert_numbered_paragraphs(folder_path)
