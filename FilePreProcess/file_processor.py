import os
import logging
import re
from docx import Document
from tqdm import tqdm

from .title_preprocess import check_lines_and_prepend
from .utils import clean_folder_name, load_module_titles
from .text_splitter import split_text_by_modules

def sanitize_filename(name: str) -> str:
    """
    过滤文件名中的非法字符，防止文件创建失败。
    Windows 不允许的字符包括：\/:*?"<>|
    """
    return re.sub(r'[\\/:\*\?"<>|]', '_', name)

def process_word_file(word_path: str, output_root: str, module_config_file: str):
    """
    读取Word文件，拆分模块，保存为txt文件。
    高鲁棒性：捕获异常，日志记录，确保目录创建。
    """
    logging.info(f"开始处理文件: {word_path}")

    if not os.path.isfile(word_path):
        logging.error(f"文件不存在或不是文件: {word_path}")
        return

    try:
        doc = Document(word_path)
    except Exception as e:
        logging.error(f"读取Word文件失败: {word_path}，错误: {e}")
        return

    folder_name = "Unnamed_Project"
    if doc.paragraphs and doc.paragraphs[0].text.strip():
        folder_name = clean_folder_name(doc.paragraphs[0].text.strip())

    output_dir = os.path.join(output_root, folder_name)
    try:
        os.makedirs(output_dir, exist_ok=True)
    except Exception as e:
        logging.error(f"创建输出目录失败: {output_dir}，错误: {e}")
        return

    full_text = "\n".join(p.text for p in doc.paragraphs)

    modules = split_text_by_modules(full_text, load_module_titles(module_config_file))
    if not modules:
        logging.warning(f"未检测到模块标题，原始文本长度：{len(full_text)}")
        complete_file = os.path.join(output_dir, "完整文本.txt")
        try:
            with open(complete_file, 'w', encoding='utf-8') as f:
                f.write(full_text)
            logging.info(f"保存完整文本文件: {complete_file}")
        except Exception as e:
            logging.error(f"保存完整文本失败，错误: {e}")
        return

    # 用tqdm包装模块迭代，显示进度条
    total = len(modules)
    with tqdm(total=total, unit="模块") as pbar:
        for title, content in modules.items():
            # 设置进度条左侧描述为当前文件编号，末尾加冒号和空格使显示更明显
            pbar.set_description(f"保存文件模块", refresh=True)
            
            try:
                if '.' in title:
                    num, mod_name = title.split('.', 1)
                    num = num.strip()
                    mod_name = mod_name.strip()
                    filename = f"{num}_{mod_name}.txt"
                else:
                    mod_name = title.strip()
                    filename = f"{mod_name}.txt"

                filename = sanitize_filename(filename)
                file_path = os.path.join(output_dir, filename)

                # 检查是否需要在文件开头添加模块编号和名称
                content = check_lines_and_prepend(content, num, mod_name)

                # 检查是否需要在文件开头添加模块编号和名称
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)

                logging.info(f"保存模块文件: {file_path}")

            except Exception as e:
                logging.error(f"保存模块文件失败: {title}，错误: {e}")

            pbar.update(1)